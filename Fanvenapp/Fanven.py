# /usr/bin/env python
# -*- coding: utf-8 -*-
# @Date    : 2018-04-27 09:49:44
# @Author  : HannochTao (hannochtao@163.com)
# @Link    : http://www.imstudy.online
# @Version : 1.4

from Translate.baidutranslate import baidu_translate, __version__
from Translate.googletranlate import google_translate
from __init__ import __author__, __copyright_info__, __doc__,__help1__,__help2__,__help3__
from tkinter import *
from tkinter.ttk import Combobox
from tkinter.scrolledtext import ScrolledText
from tkinter import messagebox
import  tkinter.filedialog
import docx
import re

class JiroApplication(Frame):


    def __init__(self, master=None):
        Frame.__init__(self, master=master)
        self.pack()
        self.create_window()

    def create_window(self):

        def translate_link(documnt_conent):
            if translate_link_chosen.get() == '百度翻译' :

                translate_result = baidu_translate(text=documnt_conent,to=language.get())
                print("baidu tranlate loading...")
                # 在翻译结果区域显示结果
                trans_result_box.insert(END, translate_result)
            elif translate_link_chosen.get() == '谷歌翻译' :
                translate_result = google_translate(text=documnt_conent,
                                                    to_language=language.get())
                print("google tranlate loading...")
                # 在翻译结果区域显示结果
                trans_result_box.insert(END, translate_result)

        def documnt_translate(event):
      
            documnt_conent = type_content_box.get('1.0', 'end-1c')
          
            # 统计有多少字符
            # print(type_content_box.count('1.0', END))
            # 匹配五段，以换行为界限
            pattern = re.compile(r'.*?\n.*?\n.*?\n.*?\n.*?\n')
            find_paragraphs = pattern.findall(documnt_conent)
            for find_paragraph in find_paragraphs:
                if type_content_box.count('1.0', END) > tuple([0]):
                	                   
                    translate_link(find_paragraph)
           
            finally_pattern = re.compile(find_paragraphs[-1] + '(.*?)')
            finally_paragraph = finally_pattern.search(documnt_conent)
            translate_link(find_paragraph)

        def boxconnent_translate():
            documnt_conent = type_content_box.get('1.0', 'end-1c')
            if type_content_box.count('1.0', END) > tuple([0]):
                translate_link(documnt_conent)
                pass



        # 显示翻译结果
        def show_trans_result():
            # 运行结果按钮状态:正在翻译
      
            trans_result_box.delete(1.0, END)
           
            run_after.config(text='正在翻译', bg='#ffffff', fg='#000000', )
            # 翻译文档 事件绑定鼠标左键按下放开时，翻译文档
            openfile_button.bind('ButtonRelease-1',documnt_translate)
            #直接翻译文本框里的内容
            
            boxconnent_translate()
            # 运行结果按钮状态:运行结果
            run_after.config(text='运行结果')

        def get_documnt(file_path):
            # 获取文档对象

            file_path_ = file_path
            file = docx.Document(file_path_)
            #print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段
            # 输出段落编号及段落内容
            type_content_box.delete(1.0,END)
            for i in range(len(file.paragraphs)):
                type_content_box.insert(END, file.paragraphs[i].text)
                type_content_box.insert(END,'\n')


        # 选择文件路径
        def selectPath():
            path_ = askdirectory()
            path.set(path_)

        def file_path():
            filetypes = [
                ("All Files", '*'),
                ("word 97-2003文档", '*.doc', 'TEXT'),
                ("word文档",'*.docx','TEXT')]

            filename = tkinter.filedialog.askopenfilename(filetypes=filetypes)
            if filename != '':
                openfile_lable.config(text=filename)
                get_documnt(filename)
            else:
                openfile_lable.config(text="您没有选择任何文件")

        def save_file():
            text_value = trans_result_box.get('1.0', END).strip()
            if text_value:
                filename = tkinter.filedialog.asksaveasfile()
                #print(filename.name)
                if filename:
                    filename.write(text_value + '\n')

                savefile_lable.config(text=filename.name)

        #打开文件标签
        openfile_lable = Label(self.master, text='请选择word文档', width=47, bg='#6495ED', fg='#ffffff',
                                  relief='flat',font=('微软雅黑', 11))
        openfile_lable.place(x=1, y=10)
        #保存文件标签
        savefile_lable = Label(self.master, text='请选择保存路径', width=40, bg='#6495ED', fg='#ffffff',
                             relief='flat', font=('微软雅黑', 11))
        savefile_lable.place(x=510, y=10)
        # 打开文件按钮
        openfile_button = Button(self.master, text="打开文件", command=lambda: file_path())
        openfile_button.place(x=430, y=10)
        #保存按钮
        save_button = Button(self.master, text="保存文件", command=lambda: save_file())
        save_button.place(x=960-80, y=10)

        ###############################################
        # 自动检测语言
        auto_language = Button(self.master, text='    自动检测语言    ', state='disabled', relief='flat',
                               bg='#ffffff', fg='#000000', font=('微软雅黑', 9), )
        auto_language.place(x=1, y=50)

        # 选择语言
        language = StringVar()
        language_chosen = Combobox(self.master, width=10, textvariable=language, )
        #language_chosen = Combobox(self.master, width=10, )
        language_chosen['values'] = ['中文', '英语', '日语']  # 设置下拉列表的值
        language_chosen.place(x=25 * 7, y=50)
        language_chosen.current(0)

        # 选择百度或者google翻译，默认为百度翻译
        translate_link_chosen = Combobox(self.master, width=10)
        translate_link_chosen['values'] = ['百度翻译', '谷歌翻译']  # 设置下拉列表的值
        translate_link_chosen.place(x=40 * 7, y=50)
        translate_link_chosen.current(0)

        # 翻译按钮
        translate_button = Button(self.master, text='翻   译', width=10, bg='#6495ED', fg='#ffffff',
                                  relief='flat', command=lambda: show_trans_result(),
                                  font=('微软雅黑', 9))
        translate_button.place(x=60 * 7 + 5, y=50)

        # 运行结果按钮
        run_after = Button(self.master, state=DISABLED, width=10, relief='flat', font=('微软雅黑', 9), )
        run_after.place(x=72 * 7 + 4, y=50)

        #########################################################################
        # 在窗口中创建输入子窗口
        type_content_box = ScrolledText(self.master, width=55, height=15, font=('微软雅黑', 11), )
        type_content_box.insert(END, '请在此处输入文字')
        type_content_box.focus_set()
        type_content_box.pack(side=LEFT, expand=True)

        # 结果窗口
        trans_result_box = ScrolledText(self.master, width=55, height=15, font=('consolas', 11), )
        trans_result_box.pack(side=RIGHT)

        ################################################
        #帮助
        help1 = Label(self.master, text=__help1__, font=('微软雅黑', 11), )
        help1.place(x=1, y = 480 - 90)

        help2 = Label(self.master, text=__help2__, font=('微软雅黑', 11), )
        help2.place(x=1, y=480 - 70)

        help3 = Label(self.master, text=__help3__, font=('微软雅黑', 11), )
        help3.place(x=1, y=480 - 50)

        # 底部版权信息
        copyright = Label(self.master, text=__copyright_info__, font=('consolas', 11), )
        copyright.place(x=(960 - len(__copyright_info__) * 7.5) / 2,
                        y=480 - 22)


WIDTH = 960
HEIGHT = 480

app_name = 'Fanven ' + __version__

root = Tk()

# 隐藏窗口
root.withdraw()

# 窗口标题
root.title(app_name)

# 设置窗口透明度 第二个参数数值越小则越透明
root.attributes('-alpha', 0.90)

# 窗口大小、居中
root.geometry('{}x{}+{}+{}'.format(str(WIDTH), str(HEIGHT),
                                   str(int((root.winfo_screenwidth() - WIDTH) / 2)),
                                   str(int((root.winfo_screenheight() - HEIGHT) / 2))))

# 固定窗口大小
root.resizable(0, 0)

# 弹出对话框
messagebox.showinfo(title='Fanven Box', message=__doc__ + '\n\t\t\t————' + __author__)
# 显示窗口
root.deiconify()

app = JiroApplication(root)

app.mainloop()
